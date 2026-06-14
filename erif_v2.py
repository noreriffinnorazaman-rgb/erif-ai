import streamlit as st
import openai
import fitz  # PyMuPDF
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import requests
import os
import json
import subprocess
import tempfile
import urllib.parse
import time

# === KONFIGURASI ===
st.set_page_config(page_title="🧠 ERIF AI v2.0 Pro", layout="wide")

# === INISIALISASI SESSION STATE ===
if "messages" not in st.session_state:
    st.session_state.messages = []
if "pdf_text" not in st.session_state:
    st.session_state.pdf_text = ""
if "api_key" not in st.session_state:
    st.session_state.api_key = "SK-GANTI-DENGAN-KEY-KAU"

# === FUNGSI EXTRACT FILE ===
def extract_pdf_text(file_path):
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        return f"Error PDF: {str(e)}"

def extract_docx_text(file_path):
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error DOCX: {str(e)}"

def extract_txt_text(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        return f"Error TXT: {str(e)}"

# === FUNGSI ROUTER AI ===
def router_decide(user_input, api_key, context=""):
    try:
        client = openai.OpenAI(api_key=api_key, base_url="https://api.moonshot.cn/v1")
        
        system_prompt = """Kau adalah Router ERIF AI. Baca soalan user dan balas JSON sahaja.

FORMAT WAJIB JSON:
{"action": "chat" atau "code" atau "generate_image" atau "export_docx" atau "export_pdf", "content": "..."}

RULES:
- "chat": Tanya konsep BM biasa. Jawab direct.
- "code": User nak kira/plot/graf/kalkulator/code Python. Content = CODE PYTHON LENGKAP.
- "generate_image": User cakap "lukis/gambar/buat image/carta/diagram". Content = prompt BM untuk image.
- "export_docx": User cakap "export word/simpan docx/save word". Content = text BM untuk Word.
- "export_pdf": User cakap "export pdf/simpan pdf/save pdf". Content = text BM untuk PDF.

Contoh 1: User "Plot graf y=x^2" → {"action": "code", "content": "import matplotlib.pyplot as plt\nimport numpy as np\nx = np.linspace(-10, 10, 100)\ny = x**2\nplt.plot(x, y)\nplt.savefig('graf.png')"}
Contoh 2: User "Lukis atom" → {"action": "generate_image", "content": "gambar atom 3D dengan elektron beredar"}
Contoh 3: User "Apa itu photosintesis" → {"action": "chat", "content": "Photosintesis adalah proses..."}"""

        messages = [
            {"role": "system", "content": system_prompt},
        ]
        
        if context:
            messages.append({"role": "system", "content": f"Context dari PDF: {context[:3000]}"})
        
        messages.append({"role": "user", "content": user_input})
        
        response = client.chat.completions.create(
            model="moonshot-v1-8k",
            messages=messages,
            temperature=0.3
        )
        
        content = response.choices[0].message.content
        # Cari JSON dalam response
        try:
            # Try parse direct
            result = json.loads(content)
        except:
            # Try extract dari markdown code block
            if "```json" in content:
                json_str = content.split("```json")[1].split("```")[0].strip()
                result = json.loads(json_str)
            elif "```" in content:
                json_str = content.split("```")[1].split("```")[0].strip()
                result = json.loads(json_str)
            else:
                # Default fallback
                result = {"action": "chat", "content": content}
        
        return result
    except Exception as e:
        return {"action": "chat", "content": f"Error router: {str(e)}"}

# === FUNGSI EXECUTOR ===
def execute_code(code, temp_dir):
    try:
        # Simpan code ke file
        code_path = os.path.join(temp_dir, "erif_code.py")
        with open(code_path, "w", encoding="utf-8") as f:
            f.write(code)
        
        # Run code
        result = subprocess.run(
            ["python", code_path],
            cwd=temp_dir,
            capture_output=True,
            text=True,
            timeout=45
        )
        
        output = result.stdout
        error = result.stderr
        
        # Cari gambar dalam temp folder
        images = []
        for f in os.listdir(temp_dir):
            if f.lower().endswith(('.png', '.jpg', '.jpeg')):
                images.append(os.path.join(temp_dir, f))
        
        return {"success": result.returncode == 0, "output": output, "error": error, "images": images}
    except subprocess.TimeoutExpired:
        return {"success": False, "output": "", "error": "Timeout: Kod ambil masa > 45 saat", "images": []}
    except Exception as e:
        return {"success": False, "output": "", "error": str(e), "images": []}

def generate_image(prompt, temp_dir):
    try:
        encoded_prompt = urllib.parse.quote(prompt)
        url = f"https://image.pollinations.ai/prompt/{encoded_prompt}?width=1024&height=1024&seed=42&nologo=true"
        
        response = requests.get(url, timeout=60)
        
        if response.status_code == 200:
            img_path = os.path.join(temp_dir, "erif_image.png")
            with open(img_path, "wb") as f:
                f.write(response.content)
            return {"success": True, "path": img_path}
        else:
            return {"success": False, "error": f"HTTP {response.status_code}"}
    except Exception as e:
        return {"success": False, "error": str(e)}

def export_docx(content, temp_dir):
    try:
        doc = Document()
        doc.add_paragraph(content)
        doc_path = os.path.join(temp_dir, "erif_export.docx")
        doc.save(doc_path)
        return {"success": True, "path": doc_path}
    except Exception as e:
        return {"success": False, "error": str(e)}

def export_pdf(content, temp_dir):
    try:
        pdf_path = os.path.join(temp_dir, "erif_export.pdf")
        doc = SimpleDocTemplate(pdf_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = [Paragraph(content.replace('\n', '<br/>'), styles["Normal"])]
        doc.build(story)
        return {"success": True, "path": pdf_path}
    except Exception as e:
        return {"success": False, "error": str(e)}

def summarize_result(action, result_content, api_key):
    try:
        client = openai.OpenAI(api_key=api_key, base_url="https://api.moonshot.cn/v1")
        
        prompt = f"Ringkaskan hasil {action} dalam BM untuk budak Tahun 5 dalam 3 ayat:\n\n{result_content}"
        
        response = client.chat.completions.create(
            model="moonshot-v1-8k",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7
        )
        
        return response.choices[0].message.content
    except:
        return "Tidak dapat ringkaskan hasil."

# === UI LAYOUT ===
st.title("🧠 ERIF AI v2.0 Pro")
st.caption("AI Tutor STEM untuk Pelajar Malaysia - Powered by Kimi API")

# Layout 2 column: Kiri Chat (70%), Kanan Sidebar (30%)
chat_col, sidebar_col = st.columns([0.7, 0.3])

# === SIDEBAR ===
with sidebar_col:
    st.header("⚙️ Tetapan & Tools")
    
    # API Key
    api_key = st.text_input(
        "🔑 Kimi API Key",
        type="password",
        value=st.session_state.api_key,
        key="api_input"
    )
    if api_key:
        st.session_state.api_key = api_key
    
    st.divider()
    
    # File Uploader
    st.subheader("📁 Upload Fail")
    uploaded_file = st.file_uploader("PDF / DOCX / TXT", type=["pdf", "docx", "txt"])
    
    if uploaded_file:
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, uploaded_file.name)
        
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        if uploaded_file.name.endswith(".pdf"):
            st.session_state.pdf_text = extract_pdf_text(file_path)
        elif uploaded_file.name.endswith(".docx"):
            st.session_state.pdf_text = extract_docx_text(file_path)
        else:
            st.session_state.pdf_text = extract_txt_text(file_path)
        
        st.success(f"✅ {uploaded_file.name} dimuatnaik!")
        st.info(f"📊 {len(st.session_state.pdf_text)} aksara diekstrak")
    
    # Clear Memory
    if st.button("🗑️ Clear Memory", use_container_width=True):
        st.session_state.messages = []
        st.session_state.pdf_text = ""
        st.success("Memory dibersihkan!")
        st.rerun()
    
    st.divider()
    
    # Info
    st.markdown("""
    **🎯 Aksi ERIF:**
    - 💬 Chat biasa
    - 🐍 Kod Python
    - 🎨 Gambar AI
    - 📄 Export DOCX
    - 📑 Export PDF
    """)

# === CHAT AREA ===
with chat_col:
    # Papar mesej sebelumnya
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            
            # Papar gambar jika ada
            if "images" in msg and msg["images"]:
                for img in msg["images"]:
                    st.image(img)
            
            # Download button jika ada
            if "file_path" in msg and msg["file_path"]:
                with open(msg["file_path"], "rb") as f:
                    file_ext = os.path.splitext(msg["file_path"])[1]
                    mime = "application/pdf" if file_ext == ".pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    st.download_button(
                        label=f"⬇️ Download {file_ext.upper()}",
                        data=f,
                        file_name=f"erif_export{file_ext}",
                        mime=mime,
                        key=f"dl_{time.time()}"
                    )
    
    # Input user
    if user_input := st.chat_input("Tanya ERIF AI sesuatu..."):
        # Tambah mesej user
        st.session_state.messages.append({"role": "user", "content": user_input})
        
        with st.chat_message("user"):
            st.markdown(user_input)
        
        # Proses dengan AI
        with st.chat_message("assistant"):
            with st.spinner("🧠 ERIF sedang berfikir..."):
                api_key = st.session_state.api_key
                context = st.session_state.pdf_text[:3000] if st.session_state.pdf_text else ""
                
                # Router decide
                router_result = router_decide(user_input, api_key, context)
                action = router_result.get("action", "chat")
                content = router_result.get("content", "")
                
                # Executor
                temp_dir = tempfile.mkdtemp()
                response_data = {"content": "", "images": [], "file_path": None}
                
                try:
                    if action == "chat":
                        response_data["content"] = content
                        st.markdown(content)
                    
                    elif action == "code":
                        st.markdown("🐍 **Menjalankan kod Python...**")
                        exec_result = execute_code(content, temp_dir)
                        
                        if exec_result["success"]:
                            output_text = exec_result["output"]
                            if exec_result["images"]:
                                output_text += f"\n\n📸 Gambar dijana: {len(exec_result['images'])} fail"
                            
                            response_data["content"] = f"```python\n{content}\n```\n\n**Output:**\n```\n{exec_result['output']}\n```"
                            response_data["images"] = exec_result["images"]
                            
                            st.markdown("✅ **Kod berjaya dijalankan!**")
                            st.code(content, language="python")
                            st.text(exec_result["output"])
                            
                            if exec_result["error"]:
                                st.warning(f"⚠️ Stderr: {exec_result['error']}")
                            
                            # Papar gambar
                            for img_path in exec_result["images"]:
                                st.image(img_path, caption="Graf/Visual dijana")
                                # Cleanup lepas 10s
                                time.sleep(1)
                        else:
                            error_msg = f"❌ **Error:** {exec_result['error']}"
                            response_data["content"] = error_msg
                            st.error(error_msg)
                    
                    elif action == "generate_image":
                        st.markdown(f"🎨 **Menjana gambar:** {content}")
                        img_result = generate_image(content, temp_dir)
                        
                        if img_result["success"]:
                            response_data["content"] = f"Gambar AI: {content}"
                            response_data["images"] = [img_result["path"]]
                            
                            st.image(img_result["path"], caption=f"Prompt: {content}")
                            
                            # Download button
                            with open(img_result["path"], "rb") as f:
                                st.download_button(
                                    "⬇️ Download Gambar",
                                    data=f,
                                    file_name="erif_image.png",
                                    mime="image/png"
                                )
                        else:
                            error_msg = f"❌ Gagal janakan gambar: {img_result['error']}"
                            response_data["content"] = error_msg
                            st.error(error_msg)
                    
                    elif action == "export_docx":
                        st.markdown("📄 **Mengeksport ke Word...**")
                        docx_result = export_docx(content, temp_dir)
                        
                        if docx_result["success"]:
                            response_data["content"] = f"📄 Dokumen Word dijana:\n\n{content[:500]}..."
                            response_data["file_path"] = docx_result["path"]
                            
                            st.success("✅ Dokumen Word sedia!")
                            st.markdown(f"**Kandungan:**\n{content[:500]}...")
                            
                            with open(docx_result["path"], "rb") as f:
                                st.download_button(
                                    "⬇️ Download DOCX",
                                    data=f,
                                    file_name="erif_export.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                        else:
                            error_msg = f"❌ Gagal export DOCX: {docx_result['error']}"
                            response_data["content"] = error_msg
                            st.error(error_msg)
                    
                    elif action == "export_pdf":
                        st.markdown("📑 **Mengeksport ke PDF...**")
                        pdf_result = export_pdf(content, temp_dir)
                        
                        if pdf_result["success"]:
                            response_data["content"] = f"📑 Dokumen PDF dijana:\n\n{content[:500]}..."
                            response_data["file_path"] = pdf_result["path"]
                            
                            st.success("✅ Dokumen PDF sedia!")
                            st.markdown(f"**Kandungan:**\n{content[:500]}...")
                            
                            with open(pdf_result["path"], "rb") as f:
                                st.download_button(
                                    "⬇️ Download PDF",
                                    data=f,
                                    file_name="erif_export.pdf",
                                    mime="application/pdf"
                                )
                        else:
                            error_msg = f"❌ Gagal export PDF: {pdf_result['error']}"
                            response_data["content"] = error_msg
                            st.error(error_msg)
                    
                    else:
                        response_data["content"] = content
                        st.markdown(content)
                    
                    # Ringkaskan hasil
                    if action in ["code", "generate_image", "export_docx", "export_pdf"]:
                        summary = summarize_result(action, response_data["content"], api_key)
                        st.divider()
                        st.info(f"📚 **Ringkasan untuk Tahun 5:**\n{summary}")
                        response_data["content"] += f"\n\n---\n📚 **Ringkasan:** {summary}"
                
                except Exception as e:
                    error_display = f"🚨 **ERIF AI error:** {str(e)}"
                    response_data["content"] = error_display
                    st.error(error_display)
                
                # Simpan ke session
                st.session_state.messages.append({
                    "role": "assistant",
                    "content": response_data["content"],
                    "images": response_data["images"],
                    "file_path": response_data["file_path"]
                })
                
                # Cleanup temp files lepas 10 saat
                time.sleep(1)
                try:
                    if os.path.exists(temp_dir):
                        for f in os.listdir(temp_dir):
                            if f.startswith("erif_"):
                                try:
                                    os.remove(os.path.join(temp_dir, f))
                                except:
                                    pass
                except:
                    pass
