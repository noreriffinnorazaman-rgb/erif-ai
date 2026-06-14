import streamlit as st
import openai
import fitz
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import requests
import os
import json
import subprocess
import tempfile
import urllib.parse
import time
import re

# === KONFIGURASI ===
st.set_page_config(page_title="ERIF AI v3.0 Ultimate", layout="wide", page_icon="🧠")

# === SESSION STATE ===
def init_session():
    defaults = {
        "messages": [], "pdf_text": "", "api_key": "",
        "use_ollama": False, "ollama_model": "llama3.2",
        "use_fooocus": False, "fooocus_url": "http://127.0.0.1:7865",
        "last_audio_path": ""
    }
    for key, val in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val
init_session()

# === PLAUSIBLE ANALYTICS ===
PLAUSIBLE_DOMAIN = os.getenv("PLAUSIBLE_DOMAIN", "")
PLAUSIBLE_SCRIPT = os.getenv("PLAUSIBLE_SCRIPT", "https://plausible.io/js/script.js")
if PLAUSIBLE_DOMAIN:
    st.components.v1.html(f'<script defer data-domain="{PLAUSIBLE_DOMAIN}" src="{PLAUSIBLE_SCRIPT}"></script>', height=0)

# === FUNGSI FILE ===
def extract_pdf_text(fp):
    try:
        doc = fitz.open(fp)
        text = "".join([p.get_text() for p in doc])
        doc.close()
        return text
    except Exception as e:
        return f"Error PDF: {e}"

def extract_docx_text(fp):
    try:
        return "\n".join([p.text for p in Document(fp).paragraphs])
    except Exception as e:
        return f"Error DOCX: {e}"

def extract_txt_text(fp):
    try:
        with open(fp, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        return f"Error TXT: {e}"

# === FUNGSI AI ===
def ollama_chat(msgs, model=None):
    try:
        model = model or st.session_state.ollama_model
        r = requests.post("http://localhost:11434/api/chat",
                       json={"model": model, "messages": msgs, "stream": False},
                       timeout=120)
        return r.json().get("message", {}).get("content", "") if r.status_code == 200 else f"Error: HTTP {r.status_code}"
    except requests.exceptions.ConnectionError:
        return "Error: Ollama tidak aktif. Jalankan 'ollama serve'"
    except Exception as e:
        return f"Error Ollama: {e}"

def ollama_list_models():
    try:
        r = requests.get("http://localhost:11434/api/tags", timeout=5)
        return [m["name"] for m in r.json().get("models", [])] if r.status_code == 200 else []
    except:
        return []

def ai_chat(msgs, api_key=None, model="moonshot-v1-8k"):
    if st.session_state.use_ollama:
        return ollama_chat(msgs)
    if not api_key and not st.session_state.api_key:
        return "Error: Tiada API Key atau Ollama"
    try:
        client = openai.OpenAI(api_key=api_key or st.session_state.api_key, base_url="https://api.moonshot.cn/v1")
        return client.chat.completions.create(model=model, messages=msgs, temperature=0.7).choices[0].message.content
    except Exception as e:
        return f"Error AI: {e}"

# === ROUTER ===
def router_decide(user_input, api_key=None, context=""):
    try:
        system = """Router ERIF v3.0 - Balas JSON: {"action": "ACTION", "content": "..."}

ACTIONS: chat | code | generate_image | generate_image_fooocus | download_video | transcribe_audio | export_docx | export_pdf

RULES:
- "download video/youtube/tiktok/ig": → download_video, content=URL
- "transcribe/transkrip/audio": → transcribe_audio
- "gambar/lukis/draw": → generate_image_fooocus (kalau Fooocus aktif) atau generate_image
- "plot/graf/python": → code
- "export word": → export_docx | "export pdf": → export_pdf
- Soalan biasa: → chat

CONTOH: {"action": "download_video", "content": "https://youtube.com/watch?v=xyz"}"""
        
        msgs = [{"role": "system", "content": system}]
        if context:
            msgs.append({"role": "system", "content": f"Context: {context[:3000]}"})
        msgs.append({"role": "user", "content": user_input})
        
        resp = ai_chat(msgs, api_key, "moonshot-v1-8k")
        if resp.startswith("Error"):
            return {"action": "chat", "content": resp}
        
        # Parse JSON
        try:
            return json.loads(resp)
        except:
            if "```json" in resp:
                return json.loads(resp.split("```json")[1].split("```")[0].strip())
            elif "```" in resp:
                return json.loads(resp.split("```")[1].split("```")[0].strip())
            else:
                m = re.search(r'\{[^}]+"action"[^}]+\}', resp, re.DOTALL)
                return json.loads(m.group()) if m else {"action": "chat", "content": resp}
    except Exception as e:
        return {"action": "chat", "content": f"Router error: {e}"}

# === TOOLS ===

def download_video(url, temp_dir, quality="best"):
    """yt-dlp downloader (FREE)"""
    try:
        tmpl = os.path.join(temp_dir, "%(title)s.%(ext)s")
        r = subprocess.run(["yt-dlp", "-f", quality, "-o", tmpl, "--no-playlist", url],
                        capture_output=True, text=True, timeout=300)
        if r.returncode == 0:
            files = [f for f in os.listdir(temp_dir) if not f.startswith("erif_") and not f.endswith(".py")]
            if files:
                fp = os.path.join(temp_dir, files[0])
                return {"success": True, "path": fp, "filename": files[0]}
        return {"success": False, "error": r.stderr or "Download failed"}
    except FileNotFoundError:
        return {"success": False, "error": "yt-dlp tidak dijumpai. Install: pip install yt-dlp"}
    except subprocess.TimeoutExpired:
        return {"success": False, "error": "Timeout > 5 minit"}
    except Exception as e:
        return {"success": False, "error": str(e)}

def get_video_info(url):
    try:
        r = subprocess.run(["yt-dlp", "-j", "--no-playlist", url], capture_output=True, text=True, timeout=30)
        if r.returncode == 0:
            info = json.loads(r.stdout)
            return {"title": info.get("title", ""), "duration": info.get("duration", 0),
                    "uploader": info.get("uploader", ""), "thumbnail": info.get("thumbnail", "")}
    except:
        pass
    return None

def transcribe_audio(fp, temp_dir, model="base"):
    """Whisper transcription (FREE)"""
    try:
        # Cuba guna whisper sebagai package
        try:
            import whisper
            m = whisper.load_model(model)
            result = m.transcribe(fp, language="ms")
            txt_path = os.path.join(temp_dir, "transcript.txt")
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(result["text"])
            return {"success": True, "text": result["text"], "path": txt_path}
        except ImportError:
            pass
        
        # Cuba subprocess
        r = subprocess.run(["whisper", fp, "--model", model, "--language", "ms",
                           "--output_dir", temp_dir, "--output_format", "txt"],
                          capture_output=True, text=True, timeout=300)
        if r.returncode == 0:
            bn = os.path.splitext(os.path.basename(fp))[0]
            txt = os.path.join(temp_dir, f"{bn}.txt")
            if os.path.exists(txt):
                with open(txt, "r", encoding="utf-8") as f:
                    return {"success": True, "text": f.read(), "path": txt}
        return {"success": False, "error": r.stderr or "Transcription failed"}
    except FileNotFoundError:
        return {"success": False, "error": "Whisper tidak dijumpai. Install: pip install openai-whisper"}
    except Exception as e:
        return {"success": False, "error": str(e)}

def check_fooocus():
    try:
        r = requests.get(f"{st.session_state.fooocus_url}/", timeout=3)
        return r.status_code == 200
    except:
        return False

def generate_image_pollinations(prompt, temp_dir):
    try:
        url = f"https://image.pollinations.ai/prompt/{urllib.parse.quote(prompt)}?width=1024&height=1024&seed=42&nologo=true"
        r = requests.get(url, timeout=120)
        if r.status_code == 200:
            fp = os.path.join(temp_dir, "erif_image.png")
            with open(fp, "wb") as f:
                f.write(r.content)
            return {"success": True, "path": fp}
        return {"success": False, "error": f"HTTP {r.status_code}"}
    except Exception as e:
        return {"success": False, "error": str(e)}

def generate_image_fooocus(prompt, temp_dir):
    """Fooocus local image generation (FREE unlimited)"""
    try:
        # Fooocus Gradio API
        api_url = f"{st.session_state.fooocus_url}/api/predict"
        payload = {
            "fn_index": 68,
            "data": [prompt, "", "cinematic-default", 1024, 1024, 1, -1, 0, "", 0.5, 1, 4,
                    "dpmpp_2m_sde_gpu", "default", 30, 0, False, 0, 1.5]
        }
        r = requests.post(api_url, json=payload, timeout=300)
        if r.status_code == 200:
            result = r.json()
            # Process response - cari image
            if "data" in result:
                for item in result["data"]:
                    if isinstance(item, str) and os.path.exists(item):
                        import shutil
                        dst = os.path.join(temp_dir, "fooocus_image.png")
                        shutil.copy(item, dst)
                        return {"success": True, "path": dst}
        return {"success": False, "error": f"HTTP {r.status_code}"}
    except requests.exceptions.ConnectionError:
        return {"success": False, "error": "Fooocus tidak aktif"}
    except Exception as e:
        return {"success": False, "error": str(e)}

def execute_code(code, temp_dir):
    try:
        cp = os.path.join(temp_dir, "code.py")
        with open(cp, "w", encoding="utf-8") as f:
            f.write(code)
        r = subprocess.run(["python", cp], cwd=temp_dir, capture_output=True, text=True, timeout=45)
        imgs = [os.path.join(temp_dir, f) for f in os.listdir(temp_dir)
                if f.lower().endswith((".png", ".jpg", ".jpeg"))]
        return {"success": r.returncode == 0, "output": r.stdout, "error": r.stderr, "images": imgs}
    except subprocess.TimeoutExpired:
        return {"success": False, "error": "Timeout 45s", "images": [], "output": ""}
    except Exception as e:
        return {"success": False, "error": str(e), "images": [], "output": ""}

def export_docx(content, temp_dir):
    try:
        d = Document()
        d.add_paragraph(content)
        fp = os.path.join(temp_dir, "export.docx")
        d.save(fp)
        return {"success": True, "path": fp}
    except Exception as e:
        return {"success": False, "error": str(e)}

def export_pdf(content, temp_dir):
    try:
        fp = os.path.join(temp_dir, "export.pdf")
        doc = SimpleDocTemplate(fp, pagesize=A4)
        doc.build([Paragraph(content.replace("\n", "<br/>"), getSampleStyleSheet()["Normal"])])
        return {"success": True, "path": fp}
    except Exception as e:
        return {"success": False, "error": str(e)}

def summarize(action, content, api_key=None):
    try:
        p = f"Ringkaskan {action} dalam BM untuk Tahun 5 (3 ayat):\n{content[:800]}"
        return ai_chat([{"role": "user", "content": p}], api_key, "moonshot-v1-8k")
    except:
        return "Tiada ringkasan"

# === UI ===
st.title("🧠 ERIF AI v3.0 Ultimate")
st.caption("AI Tutor STEM Percuma | Ollama + Kimi + Fooocus + Whisper + yt-dlp + Plausible")

chat_col, side_col = st.columns([0.7, 0.3])

# === SIDEBAR ===
with side_col:
    st.header("⚙️ Tetapan")
    
    # AI Backend
    st.subheader("🤖 AI Backend")
    st.session_state.use_ollama = st.toggle("Guna Ollama (PERCUMA)", st.session_state.use_ollama)
    
    if st.session_state.use_ollama:
        st.success("✅ Mod Percuma!")
        models = ollama_list_models()
        if models:
            st.session_state.ollama_model = st.selectbox("Model:", models, index=0)
        else:
            st.warning("Ollama tidak aktif")
            st.session_state.ollama_model = st.text_input("Model:", "llama3.2")
        st.caption("Jalankan: `ollama serve`")
    else:
        st.session_state.api_key = st.text_input("Kimi API Key:", type="password", value=st.session_state.api_key)
    
    st.divider()
    
    # Fooocus
    st.subheader("🎨 Fooocus (Unlimited)")
    st.session_state.use_fooocus = st.toggle("Aktifkan", st.session_state.use_fooocus)
    if st.session_state.use_fooocus:
        st.session_state.fooocus_url = st.text_input("URL:", st.session_state.fooocus_url)
        if check_fooocus():
            st.success("✅ Fooocus aktif")
        else:
            st.warning("⚠️ Fooocus tidak dijumpai")
    
    st.divider()
    
    # Upload
    st.subheader("📁 Upload")
    up = st.file_uploader("PDF/DOCX/TXT/Audio/Video", 
                          type=["pdf", "docx", "txt", "mp3", "mp4", "wav", "m4a", "webm"])
    if up:
        td = tempfile.gettempdir()
        fp = os.path.join(td, up.name)
        with open(fp, "wb") as f:
            f.write(up.getbuffer())
        
        if up.name.endswith(".pdf"):
            st.session_state.pdf_text = extract_pdf_text(fp)
        elif up.name.endswith(".docx"):
            st.session_state.pdf_text = extract_docx_text(fp)
        elif up.name.endswith(".txt"):
            st.session_state.pdf_text = extract_txt_text(fp)
        else:
            st.session_state.last_audio_path = fp
        
        st.success(f"✅ {up.name[:30]}...")
    
    st.divider()
    
    # Video URL
    st.subheader("📺 Video Downloader")
    vurl = st.text_input("URL Video:", placeholder="youtube.com/watch?v=...")
    if vurl and st.button("Check Video", use_container_width=True):
        info = get_video_info(vurl)
        if info:
            st.success(f"✅ {info['title'][:40]}...")
            st.caption(f"⏱️ {info['duration']//60}:{info['duration']%60:02d}")
    
    if st.button("🗑️ Clear Memory", use_container_width=True):
        st.session_state.messages = []
        st.session_state.pdf_text = ""
        st.rerun()
    
    st.divider()
    
    # Status
    st.subheader("📊 Status")
    comps = []
    if st.session_state.use_ollama:
        comps.append(("🤖 Ollama", len(ollama_list_models()) > 0))
    elif st.session_state.api_key:
        comps.append(("🌐 Kimi", True))
    else:
        comps.append(("⚠️ Tiada AI", False))
    
    comps.extend([
        ("🎨 Fooocus", st.session_state.use_fooocus and check_fooocus()),
        ("📺 yt-dlp", True),
        ("🎙️ Whisper", True)
    ])
    
    for name, ok in comps:
        st.markdown(f"{'🟢' if ok else '🔴'} {name}")
    
    st.divider()
    st.markdown("""
    **🎯 Aksi:**
    - 💬 Chat (Ollama/Percuma)
    - 🐍 Python Code
    - 🎨 Image (Fooocus/Pollinations)
    - 📺 Video Download
    - 🎙️ Transcribe
    - 📄 DOCX | 📑 PDF
    """)

# === CHAT ===
with chat_col:
    # Papar mesej
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg.get("images"):
                for img in msg["images"]:
                    st.image(img)
            if msg.get("file_path") and os.path.exists(msg["file_path"]):
                try:
                    with open(msg["file_path"], "rb") as f:
                        ext = os.path.splitext(msg["file_path"])[1].lower()
                        if ext in [".mp4", ".webm", ".mov"]:
                            st.video(msg["file_path"])
                        elif ext in [".mp3", ".wav", ".m4a"]:
                            st.audio(msg["file_path"])
                        st.download_button(f"⬇️ Download {ext}", f,
                                         file_name=f"erif{ext}",
                                         key=f"dl_{time.time()}{ext}")
                except:
                    pass
    
    # Input
    if inp := st.chat_input("Tanya ERIF atau paste URL video..."):
        st.session_state.messages.append({"role": "user", "content": inp})
        with st.chat_message("user"):
            st.markdown(inp)
        
        with st.chat_message("assistant"):
            with st.spinner("🧠 Berfikir..."):
                ctx = st.session_state.pdf_text[:3000] if st.session_state.pdf_text else ""
                router = router_decide(inp, st.session_state.api_key, ctx)
                act, content = router.get("action", "chat"), router.get("content", "")
                
                td = tempfile.mkdtemp()
                resp = {"content": "", "images": [], "file_path": None}
                
                try:
                    if act == "chat":
                        resp["content"] = content
                        st.markdown(content)
                    
                    elif act == "code":
                        st.markdown("🐍 Menjalankan kod...")
                        r = execute_code(content, td)
                        if r["success"]:
                            resp["content"] = f"```python\n{content}\n```\nOutput:\n```\n{r['output']}\n```"
                            resp["images"] = r["images"]
                            st.code(content, language="python")
                            st.text(r["output"])
                            for img in r["images"]:
                                st.image(img)
                        else:
                            st.error(f"❌ {r['error']}")
                            resp["content"] = f"Error: {r['error']}"
                    
                    elif act in ["generate_image", "generate_image_fooocus"]:
                        use_f = st.session_state.use_fooocus and check_fooocus() and act == "generate_image_fooocus"
                        st.markdown(f"🎨 Menjana gambar {'(Fooocus)' if use_f else '(Pollinations)'}...")
                        r = generate_image_fooocus(content, td) if use_f else generate_image_pollinations(content, td)
                        if r["success"]:
                            resp["content"] = f"Gambar: {content}"
                            resp["images"] = [r["path"]]
                            st.image(r["path"], caption=content)
                            with open(r["path"], "rb") as f:
                                st.download_button("⬇️ Download", f, "erif_image.png", "image/png")
                        else:
                            st.error(f"❌ {r['error']}")
                            resp["content"] = f"Error: {r['error']}"
                    
                    elif act == "download_video":
                        url = content if content.startswith("http") else inp
                        st.markdown(f"📺 Mendownload video...")
                        r = download_video(url, td)
                        if r["success"]:
                            resp["content"] = f"Video downloaded: {r['filename']}"
                            resp["file_path"] = r["path"]
                            st.success(f"✅ {r['filename']}")
                            st.video(r["path"])
                            with open(r["path"], "rb") as f:
                                ext = os.path.splitext(r["path"])[1]
                                st.download_button("⬇️ Download Video", f, r["filename"], 
                                                f"video/{ext[1:] if ext else 'mp4'}")
                        else:
                            st.error(f"❌ {r['error']}")
                            resp["content"] = f"Error: {r['error']}"
                    
                    elif act == "transcribe_audio":
                        fp = getattr(st.session_state, "last_audio_path", "")
                        if fp and os.path.exists(fp):
                            st.markdown("🎙️ Mentranskrip audio...")
                            r = transcribe_audio(fp, td)
                            if r["success"]:
                                resp["content"] = f"**Transcript:**\n\n{r['text'][:2000]}"
                                resp["file_path"] = r["path"]
                                st.success("✅ Transkripsi selesai!")
                                st.markdown(r["text"][:2000])
                                with open(r["path"], "rb") as f:
                                    st.download_button("⬇️ Download TXT", f, "transcript.txt", "text/plain")
                            else:
                                st.error(f"❌ {r['error']}")
                                resp["content"] = f"Error: {r['error']}"
                        else:
                            st.info("📤 Sila upload fail audio/video terlebih dahulu")
                            resp["content"] = "Sila upload fail audio/video"
                    
                    elif act == "export_docx":
                        st.markdown("📄 Export Word...")
                        r = export_docx(content, td)
                        if r["success"]:
                            resp["content"] = f"Dokumen Word: {content[:300]}..."
                            resp["file_path"] = r["path"]
                            st.success("✅ DOCX sedia!")
                            with open(r["path"], "rb") as f:
                                st.download_button("⬇️ Download", f, "erif.docx",
                                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        else:
                            st.error(f"❌ {r['error']}")
                    
                    elif act == "export_pdf":
                        st.markdown("📑 Export PDF...")
                        r = export_pdf(content, td)
                        if r["success"]:
                            resp["content"] = f"Dokumen PDF: {content[:300]}..."
                            resp["file_path"] = r["path"]
                            st.success("✅ PDF sedia!")
                            with open(r["path"], "rb") as f:
                                st.download_button("⬇️ Download", f, "erif.pdf", "application/pdf")
                        else:
                            st.error(f"❌ {r['error']}")
                    
                    else:
                        st.markdown(content)
                        resp["content"] = content
                    
                    # Ringkasan
                    if act in ["code", "generate_image", "generate_image_fooocus", "download_video", "transcribe_audio"]:
                        s = summarize(act, resp["content"], st.session_state.api_key)
                        st.divider()
                        st.info(f"📚 Ringkasan (Tahun 5): {s}")
                        resp["content"] += f"\n\n**Ringkasan:** {s}"
                
                except Exception as e:
                    st.error(f"🚨 Error: {e}")
                    resp["content"] = f"Error: {e}"
                
                st.session_state.messages.append(resp)
                time.sleep(0.5)
                try:
                    if os.path.exists(td):
                        for f in os.listdir(td):
                            try:
                                os.remove(os.path.join(td, f))
                            except:
                                pass
                except:
                    pass
